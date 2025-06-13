"""
Document Loader Module - Handles multiple file formats and chunking
"""

import os
import io
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Multi-format document loader"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize document loader
        
        Args:
            chunk_size: Text chunk size
            chunk_overlap: Chunk overlap size
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Load PDF file"""
        try:
            doc = fitz.open(file_path)
            pages_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                if text.strip():  # Only process non-empty pages
                    pages_content.append({
                        'content': text,
                        'page_num': page_num + 1,
                        'file_name': Path(file_path).name,
                        'file_type': 'pdf'
                    })
            
            doc.close()
            return pages_content
            
        except Exception as e:
            logger.error(f"PDF loading error: {e}")
            return []
    
    def load_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Load DOCX file"""
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            full_text = '\n'.join(content)
            
            return [{
                'content': full_text,
                'page_num': 1,
                'file_name': Path(file_path).name,
                'file_type': 'docx'
            }] if full_text.strip() else []
            
        except Exception as e:
            logger.error(f"DOCX loading error: {e}")
            return []
    
    def load_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """Load TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return [{
                'content': content,
                'page_num': 1,
                'file_name': Path(file_path).name,
                'file_type': 'txt'
            }] if content.strip() else []
            
        except UnicodeDecodeError:
            # Try other encodings
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
                
                return [{
                    'content': content,
                    'page_num': 1,
                    'file_name': Path(file_path).name,
                    'file_type': 'txt'
                }] if content.strip() else []
                
            except Exception as e:
                logger.error(f"TXT loading error: {e}")
                return []
        except Exception as e:
            logger.error(f"TXT loading error: {e}")
            return []
    
    def load_from_bytes(self, file_bytes: bytes, file_name: str) -> List[Dict[str, Any]]:
        """Load file from bytes"""
        file_extension = Path(file_name).suffix.lower()
        
        # Create temporary file
        temp_path = f"/tmp/{file_name}"
        
        try:
            with open(temp_path, 'wb') as f:
                f.write(file_bytes)
            
            if file_extension == '.pdf':
                result = self.load_pdf(temp_path)
            elif file_extension == '.docx':
                result = self.load_docx(temp_path)
            elif file_extension == '.txt':
                result = self.load_txt(temp_path)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return []
            
            return result
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def chunk_documents(self, documents: List[Dict[str, Any]]) -> List[LangChainDocument]:
        """Chunk documents"""
        langchain_docs = []
        
        for doc in documents:
            # Use text splitter for chunking
            chunks = self.text_splitter.split_text(doc['content'])
            
            for i, chunk in enumerate(chunks):
                if chunk.strip():  # Only keep non-empty chunks
                    metadata = {
                        'file_name': doc['file_name'],
                        'file_type': doc['file_type'],
                        'page_num': doc['page_num'],
                        'chunk_id': i,
                        'total_chunks': len(chunks)
                    }
                    
                    langchain_docs.append(
                        LangChainDocument(
                            page_content=chunk,
                            metadata=metadata
                        )
                    )
        
        return langchain_docs
    
    def load_multiple_files(self, file_list: List[tuple]) -> List[LangChainDocument]:
        """Load multiple files
        
        Args:
            file_list: List of (file_bytes, file_name) tuples
            
        Returns:
            List of chunked LangChain documents
        """
        all_documents = []
        
        for file_bytes, file_name in file_list:
            logger.info(f"Loading file: {file_name}")
            docs = self.load_from_bytes(file_bytes, file_name)
            all_documents.extend(docs)
        
        # Chunk processing
        chunked_docs = self.chunk_documents(all_documents)
        
        logger.info(f"Generated {len(chunked_docs)} document chunks")
        return chunked_docs 